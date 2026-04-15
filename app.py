"""
EÜTF Sınav Analiz ve Rapor Sistemi v4
Ege Üniversitesi Tıp Fakültesi — Tıp Eğitimi Anabilim Dalı
BYS / SBYS uyumlu · Gelişmiş psikometri · Karar destek matrisi
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re, warnings, datetime

warnings.filterwarnings("ignore")

# ============================================================
# PAGE CONFIG
# ============================================================
st.set_page_config(
    page_title="EÜTF Sınav Analiz",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.title("🏥 EÜTF TEAD Sınav Analiz Raporu Yazılımı™")
st.caption("BYS / SBYS uyumlu · Gelişmiş psikometri · Karar destek matrisi · AI Destekli Değerlendirme")


# ============================================================
# GEMINI API KEY — önce Streamlit secrets, sonra kullanıcı input
# ============================================================
def get_gemini_key():
    """Streamlit Cloud secrets veya kullanıcı inputundan API key al."""
    try:
        return st.secrets["GEMINI_KEY"]
    except Exception:
        return None


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def read_flex(f):
    """BYS/SBYS .xls/.xlsx — format ne olursa olsun oku."""
    buf = BytesIO(f.read())
    f.seek(0)
    for eng in ["openpyxl", "xlrd", None]:
        try:
            buf.seek(0)
            kw = {"engine": eng} if eng else {}
            return pd.read_excel(buf, header=None, **kw)
        except Exception:
            continue
    st.error("Dosya okunamadı. Lütfen Excel'de açıp .xlsx olarak yeniden kaydedin.")
    st.stop()


def find_hdr(raw, markers):
    """Header satırını bul."""
    for i in range(min(15, len(raw))):
        vs = [str(v).strip().lower() for v in raw.iloc[i] if pd.notna(v)]
        if all(any(m in v for v in vs) for m in markers):
            return i
    return None


def read_hdr(f, h):
    """Dosyayı belirli header satırıyla oku."""
    f.seek(0)
    buf = BytesIO(f.read())
    for eng in ["openpyxl", "xlrd", None]:
        try:
            buf.seek(0)
            kw = {"engine": eng} if eng else {}
            return pd.read_excel(buf, header=h, **kw)
        except Exception:
            continue
    return None


def parse_student(f):
    """Öğrenci Soru Analizi dosyasını parse et."""
    raw = read_flex(f)
    h = find_hdr(raw, ["q1"])
    if h is None:
        st.error("Öğrenci dosyasında 'q1' header satırı bulunamadı.")
        st.stop()
    df = read_hdr(f, h)
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all")

    qc = sorted(
        [c for c in df.columns if re.match(r"^[qQ]\d+$", c)],
        key=lambda x: int(re.search(r"\d+", x).group()),
    )
    if not qc:
        st.error("Soru kolonları bulunamadı.")
        st.stop()

    df.rename(columns={c: c.upper() for c in qc}, inplace=True)
    qc = [c.upper() for c in qc]
    for c in qc:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).astype(int)

    df["TOTAL_RAW"] = df[qc].sum(axis=1)
    df["TOTAL"] = df["TOTAL_RAW"] / len(qc) * 100

    return df[df["TOTAL_RAW"] > 0].copy(), qc


def parse_freq(f):
    """Frekans Analizi dosyasını parse et."""
    raw = read_flex(f)
    h = find_hdr(raw, ["#", "soru"])
    if h is None:
        h = find_hdr(raw, ["zorluk"])
    if h is None:
        st.error("Frekans dosyasında header satırı bulunamadı.")
        st.stop()

    freq = read_hdr(f, h)
    cmap = {}
    for c in freq.columns:
        if pd.isna(c):
            continue
        cs = str(c).strip().lower()
        if cs == "#" or cs == "no":
            cmap[c] = "NO"
        elif "seçenek" in cs:
            cmap[c] = "SEC"
        elif "zorluk" in cs:
            cmap[c] = "ZORLUK"
        elif "ayırt" in cs:
            cmap[c] = "AYIRT"
        elif "başarı" in cs:
            cmap[c] = "BASARI"
        elif "soru kök" in cs:
            cmap[c] = "KOK"
        elif "soru sahibi" in cs:
            cmap[c] = "SAHIP"
        elif "kaynak" in cs:
            cmap[c] = "KAYNAK"
    freq.rename(columns=cmap, inplace=True)

    # Başlıksız sütunları tespit et — daha önce çıktığı sınav bilgisi
    # Genellikle son 2 sütun: sınav adı + eski zorluk değeri
    unnamed_cols = [c for c in freq.columns if str(c).startswith("Unnamed")]
    if len(unnamed_cols) >= 1:
        freq.rename(columns={unnamed_cols[0]: "ONCEKI_SINAV"}, inplace=True)
    if len(unnamed_cols) >= 2:
        freq.rename(columns={unnamed_cols[1]: "ONCEKI_P"}, inplace=True)

    if "NO" in freq.columns:
        freq = freq[
            freq["NO"].apply(
                lambda x: (
                    pd.notna(x)
                    and isinstance(x, (int, float, np.integer, np.floating))
                    and float(x) == int(float(x))
                )
                if pd.notna(x)
                else False
            )
        ].copy()
        freq["NO"] = freq["NO"].astype(int)
    return freq


def parse_dist(sec, n):
    """SBYS Seçenekler hücresinden çeldirici frekanslarını parse et."""
    if pd.isna(sec):
        return []
    parts = re.findall(
        r"\(CS:\s*(\d+)\)\s*-\s*(.+?)(?=\\n|\n|\(CS:|$)", str(sec)
    )
    return [
        {
            "count": int(c),
            "text": t.replace("*", "").strip(),
            "correct": "*" in t,
            "pct": int(c) / n * 100 if n > 0 else 0,
            "func": int(c) / n * 100 >= 5 or "*" in t,
        }
        for c, t in parts
    ]


def kr20(di):
    """KR-20 hesapla."""
    p = di.mean(0)
    tv = di.sum(1).var(ddof=1)
    n = di.shape[1]
    return (n / (n - 1)) * (1 - (p * (1 - p)).sum() / tv) if tv > 0 else 0


def d_idx(df, qc):
    """Üst-alt %27 D indeksi hesapla."""
    t = df[qc].sum(1)
    k = max(1, int(len(t) * 0.27))
    s = t.sort_values()
    lo, hi = s.index[:k], s.index[-k:]
    return {
        q: round(df.loc[hi, q].mean() - df.loc[lo, q].mean(), 4) for q in qc
    }


def cat_d(p):
    if p >= 0.85:
        return "Çok Kolay"
    if p >= 0.60:
        return "Önerilen"
    if p >= 0.30:
        return "Kabul Edilebilir"
    return "Çok Zor"


def cat_disc(d):
    if d >= 0.40:
        return "Mükemmel"
    if d >= 0.30:
        return "İyi"
    if d >= 0.20:
        return "Düzeltilmeli"
    return "Kullanılmamalı"


def karar_fn(p, d, r):
    if d < 0:
        return "Çıkar (negatif ayırt edici)"
    if 0.30 <= p <= 0.80 and d >= 0.30:
        return "Sakla"
    if d < 0.20 and r < 0.15:
        return "Çıkar"
    if d < 0.20:
        return "Revize et"
    if p >= 0.95:
        return "Çok kolay — revize"
    if p < 0.15:
        return "Çok zor — revize"
    return "Gözden geçir"


def ferguson_delta(scores, k):
    n = len(scores)
    freq = pd.Series(scores).value_counts()
    return (n**2 - ((freq**2).sum())) / (n**2 - n**2 / (k + 1))


def extract_exam_name(filename):
    """Excel dosya adından sınav adını çıkar.
    Örn: 'ÇOCUK SAĞLIĞI ... SINAVI (SBYS) - Tarih_ ...' → 'ÇOCUK SAĞLIĞI ... SINAVI (SBYS)'
    """
    name = filename
    # Uzantıyı kaldır
    name = re.sub(r'\.(xlsx?|csv)$', '', name, flags=re.IGNORECASE)
    # ' - Tarih_' veya ' - Tarih:' sonrasını kaldır
    name = re.split(r'\s*-\s*Tarih[_:]', name, flags=re.IGNORECASE)[0]
    # 'Öğrenci Soru Analizi' / 'Frekans Analizi' sonrasını kaldır
    name = re.split(r'\s*(Öğrenci Soru Analizi|Frekans Analizi)', name, flags=re.IGNORECASE)[0]
    return name.strip()


# ============================================================
# UI — INPUT
# ============================================================
st.divider()

c1, c2 = st.columns(2)
with c1:
    student_file = st.file_uploader(
        "📊 Öğrenci Soru Analizi (BYS/SBYS)", type=["xls", "xlsx"]
    )
with c2:
    freq_file = st.file_uploader(
        "📋 Frekans Analizi (BYS/SBYS)", type=["xls", "xlsx"]
    )

c3, c4 = st.columns(2)
with c3:
    analyst_name = st.text_input(
        "👤 Danışman Öğretim Üyesi (Ünvan – Ad Soyad)",
        placeholder="Doç. Dr. ...",
    )
with c4:
    secrets_key = get_gemini_key()
    if secrets_key:
        api_key = secrets_key
        st.success("🔑 Gemini API Key: secrets'tan yüklendi")
    else:
        api_key = st.text_input(
            "🔑 Gemini API Key (opsiyonel)", type="password"
        )

run_btn = st.button(
    "🚀 Analizi Başlat", type="primary", use_container_width=True
)


# ============================================================
# MAIN PIPELINE
# ============================================================
if run_btn and student_file and freq_file:

    # ---- Data Loading ----
    with st.spinner("Veriler okunuyor..."):
        df, q_cols = parse_student(student_file)
        freq = parse_freq(freq_file)

    # Sınav adını dosya adından çıkar
    exam_name = extract_exam_name(student_file.name)

    N = len(df)
    K = len(q_cols)
    scores = df["TOTAL"].values.astype(float)  # 100 üzerinden
    raw_scores = df["TOTAL_RAW"].values.astype(float)

    st.success(
        f"✅ {N} öğrenci × {K} madde yüklendi. "
        f"Puanlar 100 üzerinden normalize edildi."
    )

    # ---- Analysis ----
    with st.spinner("Psikometrik analiz yapılıyor..."):
        mn = scores.mean()
        sd = scores.std(ddof=1)
        med = np.median(scores)
        q1v = np.percentile(scores, 25)
        q3v = np.percentile(scores, 75)
        skew = pd.Series(scores).skew()
        kurt = pd.Series(scores).kurtosis()
        alpha = kr20(df[q_cols])
        sem = sd * np.sqrt(1 - alpha)
        ci95 = 1.96 * sem
        fdelta = ferguson_delta(raw_scores, K)

        # Split-half
        odd_c = [q_cols[i] for i in range(0, len(q_cols), 2)]
        even_c = [q_cols[i] for i in range(1, len(q_cols), 2)]
        r_half = df[odd_c].sum(1).corr(df[even_c].sum(1))
        sb_corr = 2 * r_half / (1 + r_half)
        guttman = 2 * (
            1
            - (df[odd_c].sum(1).var(ddof=1) + df[even_c].sum(1).var(ddof=1))
            / raw_scores.var(ddof=1)
        )

        # D-index & corrected rpbi
        dv = d_idx(df, q_cols)
        rpbi = {}
        for q in q_cols:
            r = df[q].corr(df["TOTAL_RAW"] - df[q])
            rpbi[q] = round(r, 4) if pd.notna(r) else 0

        # Item analysis
        rows = []
        for q in q_cols:
            qn_num = int(re.search(r"\d+", q).group())
            p = df[q].mean()
            d = dv[q]
            r = rpbi[q]

            fr = (
                freq[freq["NO"] == qn_num]
                if "NO" in freq.columns
                else pd.DataFrame()
            )
            nf = td = 0
            onceki_sinav = ""
            onceki_p = None
            if not fr.empty:
                if "SEC" in freq.columns:
                    for o in parse_dist(fr.iloc[0]["SEC"], N):
                        if not o["correct"]:
                            td += 1
                            nf += 0 if o["func"] else 1
                # Daha önce çıktığı sınav bilgisi (Unnamed sütundan)
                if "ONCEKI_SINAV" in freq.columns:
                    s_val = fr.iloc[0]["ONCEKI_SINAV"]
                    if pd.notna(s_val):
                        s_str = str(s_val).strip()
                        if s_str and s_str not in ("0", "0.0", "nan"):
                            onceki_sinav = s_str
                # Önceki sınavdaki zorluk değeri
                if "ONCEKI_P" in freq.columns:
                    p_val = fr.iloc[0]["ONCEKI_P"]
                    if pd.notna(p_val):
                        try:
                            onceki_p = round(float(p_val), 4)
                        except (ValueError, TypeError):
                            onceki_p = None

            ds = "—" if td == 0 else ("✅" if nf == 0 else f"{nf}/{td}")
            # Tekrar durumu etiketi
            if onceki_sinav:
                tekrar_label = f"Tekrar ({onceki_sinav})"
                if onceki_p is not None:
                    tekrar_label += f" [p={onceki_p}]"
            else:
                tekrar_label = "Yeni"

            rows.append(
                {
                    "Madde": q,
                    "p": round(p, 2),
                    "Zorluk": cat_d(p),
                    "D": round(d, 2),
                    "r_pbi": round(r, 2),
                    "Kategori": cat_disc(d),
                    "Çeldirici": ds,
                    "Önceki Kullanım": tekrar_label,
                    "Karar": karar_fn(p, d, r),
                }
            )
        item_df = pd.DataFrame(rows)
        cc = item_df["Kategori"].value_counts()

        # Quality subset KR-20
        qm = item_df["Zorluk"].isin(["Önerilen", "Kabul Edilebilir"]) & item_df[
            "Kategori"
        ].isin(["Mükemmel", "İyi"])
        qi = item_df[qm]["Madde"].tolist()
        kr20_q = kr20(df[qi]) if len(qi) > 1 else None

    # ============================================================
    # DISPLAY
    # ============================================================

    # ---- Kısaltmalar ----
    with st.expander("📖 Kısaltmalar ve Temel Kavramlar Rehberi", expanded=False):
        st.markdown(
            """
| Terim | Açıklama |
|:---|:---|
| **p (Güçlük İndeksi)** | Maddeyi doğru yanıtlayan öğrenci oranı. 0.30–0.80 ideal. |
| **D (Ayırt Edicilik İndeksi)** | Üst %27 ve alt %27 gruplarının doğru yanıt farkı. D≥0.30 iyi; D<0.20 zayıf; D<0 sorunlu. |
| **r_pbi** | Düzeltilmiş nokta çift serili korelasyon (madde–toplam). r_pbi≥0.30 iyi; <0.15 zayıf. |
| **KR-20** | Kuder-Richardson 20 — dikotom maddeler için iç tutarlılık. ≥0.80 yüksek; 0.70–0.79 kabul edilebilir. |
| **SEM** | Standart Ölçme Hatası = SD × √(1 − KR-20). Gözlenen puandaki belirsizlik bandı. |
| **%95 Güven Aralığı** | Gerçek puan ±1.96×SEM aralığında olma olasılığı %95. |
| **Ferguson's δ** | Puanların yayılım genişliği. δ>0.90 iyi ayırt etme kapasitesi. |
| **Guttman Split-Half** | Sınavın tek-çift yarıları arası güvenirlik. KR-20'ye alternatif kanıt. |
| **Çeldirici** | Yanlış seçenek. Fonksiyonel çeldirici ≥%5 oranında seçilmelidir. |
| **Negatif Ayırt Edici** | D<0: alt grup üst gruptan daha başarılı → anahtarlama hatası veya belirsiz soru. |
| **Tavan Etkisi** | p≥0.95: neredeyse herkes doğru → ayırt edicilik düşük. |
| **Taban Etkisi** | p≤0.20: neredeyse kimse doğru yapamıyor. |
"""
        )

    # ---- KTT Notu ----
    with st.expander(
        "ℹ️ Neden Klasik Test Teorisi (KTT) ve Üst-Alt Grup Yöntemi?",
        expanded=False,
    ):
        st.markdown(
            f"""
Bu rapordaki madde analizleri **Klasik Test Teorisi (KTT)** çerçevesinde, **üst-alt %27 grup yöntemiyle** yapılmıştır.

**Neden KTT?** KTT, tıp eğitiminde dünya genelinde en yaygın kullanılan psikometrik yaklaşımdır. NBME, AMEE ve TEPDAD standartları KTT metriklerini referans alır. Hesaplaması şeffaf, yorumlaması kolaydır ve öğretim üyelerine doğrudan geri bildirim verilmesine uygundur.

**Neden üst-alt %27?** Kelley (1939) ve Ebel (1965) tarafından gösterildiği üzere, toplam puan dağılımının üst ve alt %27'lik dilimleri kullanıldığında, ayırt edicilik indeksi (D) istatistiksel olarak en güçlü kestirim performansını gösterir. Bu yöntem örneklem büyüklüğünden nispeten bağımsızdır ve küçük gruplarda bile güvenilir sonuçlar verir.

**Alternatifler:** Madde Yanıt Kuramı (MYK / IRT), özellikle büyük örneklemlerde (N>500) ve soru havuzu kalibrasyonunda üstün performans gösterir. Ancak mevcut öğrenci sayısı (N={N}) ile KTT daha uygun bir yaklaşımdır.
"""
        )

    # ---- Genel Göstergeler ----
    st.header("📊 Genel Psikometrik Göstergeler")
    r1 = st.columns(4)
    r1[0].metric("Ortalama ± SD", f"{mn:.1f} ± {sd:.1f}")
    r1[1].metric("Ortanca", f"{med:.1f}")
    r1[2].metric(
        "KR-20 (tüm maddeler)",
        f"{alpha:.3f}",
        delta=(
            "Yüksek"
            if alpha >= 0.80
            else "Kabul edilebilir" if alpha >= 0.70 else "Düşük"
        ),
    )
    r1[3].metric("SEM", f"{sem:.2f}")

    r2 = st.columns(4)
    r2[0].metric(
        "Ferguson's δ",
        f"{fdelta:.3f}",
        delta="İyi" if fdelta >= 0.90 else "Düşük",
    )
    r2[1].metric("Guttman Split-Half", f"{guttman:.3f}")
    r2[2].metric("Spearman-Brown", f"{sb_corr:.3f}")
    r2[3].metric("%95 Güven Aralığı", f"±{ci95:.1f} puan")

    if kr20_q is not None:
        st.info(
            f"🎯 **KR-20 (kaliteli alt küme — {len(qi)} madde):** {kr20_q:.3f}  \n"
            f"Önerilen/Kabul Edilebilir güçlük ∩ İyi/Mükemmel ayırt edicilik: "
            f"{', '.join(qi)}"
        )
        

    r3 = st.columns(4)
    p60 = (scores >= 60).sum()
    p50 = (scores >= 50).sum()
    f40 = (scores < 40).sum()
    r3[0].metric("≥60 puan", f"{p60} (%{p60 / N * 100:.1f})")
    r3[1].metric("≥50 puan", f"{p50} (%{p50 / N * 100:.1f})")
    r3[2].metric("<40 puan", f"{f40} (%{f40 / N * 100:.1f})")
    r3[3].metric("Çarpıklık / Basıklık", f"{skew:.2f} / {kurt:.2f}")
    st.caption(f"Not: Puanlar {K} madde üzerinden 100'e normalize edilmiştir.")

    # ---- Histogram + Normal Eğri ----
    st.header("📈 Öğrencilerin Puan Dağılımı")
    fig1, ax1 = plt.subplots(figsize=(10, 4.5))
    bins = np.arange(
        int(scores.min() // 10 * 10), int(scores.max() // 10 * 10) + 20, 10
    )
    cnts, edges, patches = ax1.hist(
        scores,
        bins=bins,
        edgecolor="black",
        color="#4472C4",
        alpha=0.75,
        density=False,
        label="Gözlenen dağılım",
    )
    x_n = np.linspace(scores.min() - 10, scores.max() + 10, 200)
    y_n = sp_stats.norm.pdf(x_n, mn, sd) * N * (edges[1] - edges[0])
    ax1.plot(
        x_n,
        y_n,
        color="#C0392B",
        linewidth=2,
        ls="--",
        label=f"Normal dağılım eğrisi (μ={mn:.1f}, σ={sd:.1f})",
    )
    for cv, pv in zip(cnts, patches):
        if cv > 0:
            ax1.text(
                pv.get_x() + pv.get_width() / 2,
                cv + 0.8,
                str(int(cv)),
                ha="center",
                fontsize=9,
                fontweight="bold",
            )
    ax1.axvline(mn, color="#E67E22", ls=":", alpha=0.8, label=f"Ortalama = {mn:.1f}")
    ax1.axvline(med, color="#27AE60", ls=":", alpha=0.8, label=f"Ortanca = {med:.1f}")
    ax1.set_xlabel("Puan (100 üzerinden)")
    ax1.set_ylabel("Öğrenci Sayısı")
    ax1.legend(fontsize=8, loc="upper left")
    ax1.grid(axis="y", alpha=0.2)
    plt.tight_layout()
    st.pyplot(fig1)

    # ---- Violin Plot ----
    st.header("🎻 Öğrencilerin Puan Dağılımı — Violin Plot")
    fig_v, ax_v = plt.subplots(figsize=(10, 3.5))
    vp = ax_v.violinplot(
        [scores],
        positions=[0],
        vert=False,
        showmeans=True,
        showmedians=True,
        showextrema=True,
    )
    vp["bodies"][0].set_facecolor("#4472C4")
    vp["bodies"][0].set_alpha(0.5)
    if "cmeans" in vp:
        vp["cmeans"].set_color("#E67E22")
        vp["cmeans"].set_linewidth(2)
    if "cmedians" in vp:
        vp["cmedians"].set_color("#27AE60")
        vp["cmedians"].set_linewidth(2)
    for part in ["cmins", "cmaxes", "cbars"]:
        if part in vp:
            vp[part].set_color("#555")
    ax_v.scatter(
        scores,
        np.random.normal(0, 0.03, len(scores)),
        alpha=0.3,
        s=8,
        color="#2C3E50",
        zorder=3,
    )
    ax_v.set_xlabel("Puan (100 üzerinden)")
    ax_v.set_yticks([])
    legend_els = [
        Line2D([0], [0], color="#4472C4", lw=8, alpha=0.5, label="Yoğunluk tahmini (kernel density)"),
        Line2D([0], [0], color="#E67E22", lw=2, label=f"Ortalama ({mn:.1f})"),
        Line2D([0], [0], color="#27AE60", lw=2, label=f"Ortanca ({med:.1f})"),
        Line2D([0], [0], color="#555", lw=1, label=f"Min–Max ({scores.min():.0f}–{scores.max():.0f})"),
        Line2D([0], [0], marker="o", color="w", markerfacecolor="#2C3E50", markersize=4, alpha=0.5, label="Bireysel puanlar"),
    ]
    ax_v.legend(handles=legend_els, fontsize=7, loc="upper left", framealpha=0.9)
    ax_v.grid(axis="x", alpha=0.2)
    plt.tight_layout()
    st.pyplot(fig_v)

    # ---- Madde Analizi ----
    st.header("🔬 Madde Analizi")

    # Tekrar eden soru istatistikleri
    tekrar_df = item_df[item_df["Önceki Kullanım"] != "Yeni"]
    tekrar_count = len(tekrar_df)
    yeni_count = len(item_df) - tekrar_count

    mc = st.columns(6)
    mc[0].metric("🟢 Mükemmel", cc.get("Mükemmel", 0))
    mc[1].metric("🟡 İyi", cc.get("İyi", 0))
    mc[2].metric("🟠 Düzeltilmeli", cc.get("Düzeltilmeli", 0))
    mc[3].metric("🔴 Kullanılmamalı", cc.get("Kullanılmamalı", 0))
    mc[4].metric("⛔ Negatif Ayırt Edici", (item_df["D"] < 0).sum())
    mc[5].metric("🔄 Tekrar Eden Soru", tekrar_count)

    if tekrar_count > 0:
        tekrar_sinavlar = tekrar_df["Önceki Kullanım"].apply(
            lambda x: re.search(r"Tekrar \((.+?)\)", x).group(1) if re.search(r"Tekrar \((.+?)\)", x) else ""
        ).value_counts()
        sinav_ozet = ", ".join([f"{s}: {c} soru" for s, c in tekrar_sinavlar.items() if s])
        st.info(f"🔄 **{tekrar_count}/{K} madde daha önce kullanılmış** ({sinav_ozet})")

    display_df = item_df.rename(
        columns={"Çeldirici": "Çeldirici\n(işlevsiz/toplam)"}
    )

    def clr(v):
        m = {
            "Mükemmel": "#E8F5E9",
            "İyi": "#F1F8E9",
            "Düzeltilmeli": "#FFF3E0",
            "Kullanılmamalı": "#FFEBEE",
        }
        return f"background-color: {m[v]}" if v in m else ""

    st.dataframe(
       display_df.style.map(clr, subset=["Kategori"]),
        use_container_width=True,
        height=500,
    )

    # ---- Karar Destek Matrisi ----
    st.header("🧭 Karar Destek Matrisi")
    st.caption(
        "Aşağıdaki tabloda sınav setine ilişkin karar destek matrisi sunulmaktadır. "
        "Tablodaki renkler maddelerin ayırt edicilik ve zorluk indeksine göre dağılımını "
        "ifade etmektedir. Buna göre yeşil = sakla, sarı = gözden geçir, turuncu = revize et "
        "ve kırmızı = sınav setinden çıkar anlamına gelmektedir."
    )
    p_labs = [
        "Çok Zor\n(p<0.30)",
        "Kabul Ed.\n(0.30–0.59)",
        "Önerilen\n(0.60–0.84)",
        "Çok Kolay\n(p≥0.85)",
    ]
    d_labs = [
        "Kullanılmamalı\n(D<0.20)",
        "Düzeltilmeli\n(0.20–0.29)",
        "İyi\n(0.30–0.39)",
        "Mükemmel\n(D≥0.40)",
    ]
    mx = np.zeros((4, 4), dtype=int)
    for _, r in item_df.iterrows():
        pi = 0 if r["p"] < 0.30 else (1 if r["p"] < 0.60 else (2 if r["p"] < 0.85 else 3))
        di = 0 if r["D"] < 0.20 else (1 if r["D"] < 0.30 else (2 if r["D"] < 0.40 else 3))
        mx[di][pi] += 1

    fig_m, ax_m = plt.subplots(figsize=(9, 5))
    cm_ = np.array(
        [
            ["#FFCDD2", "#FFF9C4", "#C8E6C9", "#FFCDD2"],
            ["#FFE0B2", "#FFF9C4", "#C8E6C9", "#FFE0B2"],
            ["#C8E6C9", "#C8E6C9", "#A5D6A7", "#FFF9C4"],
            ["#A5D6A7", "#A5D6A7", "#A5D6A7", "#FFF9C4"],
        ]
    )
    for i in range(4):
        for j in range(4):
            ax_m.add_patch(
                plt.Rectangle(
                    (j, i), 1, 1, facecolor=cm_[i][j], edgecolor="white", linewidth=2
                )
            )
            ax_m.text(
                j + 0.5, i + 0.5, str(mx[i][j]),
                ha="center", va="center", fontsize=18, fontweight="bold",
            )
    ax_m.set_xlim(0, 4)
    ax_m.set_ylim(0, 4)
    ax_m.set_xticks([0.5, 1.5, 2.5, 3.5])
    ax_m.set_xticklabels(p_labs, fontsize=8)
    ax_m.set_yticks([0.5, 1.5, 2.5, 3.5])
    ax_m.set_yticklabels(d_labs, fontsize=8)
    ax_m.set_xlabel("Zorluk (p)", fontsize=10)
    ax_m.set_ylabel("Ayırt Edicilik (D)", fontsize=10)
    ax_m.set_title("Karar Destek Matrisi", fontsize=12, fontweight="bold")
    ax_m.invert_yaxis()
    plt.tight_layout()
    st.pyplot(fig_m)
    st.caption(
        "🟢 Yeşil = Sakla  |  🟡 Sarı = Gözden geçir  |  🟠 Turuncu = Revize  |  🔴 Kırmızı = Çıkar"
    )

    # ---- Scatter ----
    st.header("📉 Zorluk ve Ayırt Edicilik İndekslerinin Birlikteliği")
    fig2, ax2 = plt.subplots(figsize=(10, 6))
    cm2 = {
        "Mükemmel": "#2E7D32",
        "İyi": "#4CAF50",
        "Düzeltilmeli": "#FF9800",
        "Kullanılmamalı": "#F44336",
    }
    ax2.scatter(
        item_df["p"], item_df["D"],
        c=[cm2.get(r["Kategori"], "#999") for _, r in item_df.iterrows()],
        s=60, alpha=0.8, edgecolors="black", linewidths=0.5,
    )
    ax2.axhline(y=0.20, color="red", ls="--", alpha=0.5)
    ax2.axvline(x=0.30, color="blue", ls="--", alpha=0.3)
    ax2.axvline(x=0.80, color="blue", ls="--", alpha=0.3)
    ax2.fill_betweenx([0.20, 0.70], 0.30, 0.80, alpha=0.06, color="green")
    ax2.set_xlabel("Zorluk İndeksi (p)")
    ax2.set_ylabel("Ayırt Edicilik İndeksi (D)")
    ax2.grid(alpha=0.2)
    plt.tight_layout()
    st.pyplot(fig2)

    # ---- Kesme Puanı Simülasyonu ----
    st.header("🔮 Kesme Puanı Simülasyonu")
    st.caption(
        "Kesme puanı simülasyonu; sınav analizleri sonrasında sınav setindeki kullanılmamalı "
        "kategorisindeki sorular dikkate alınmadan analizler yeniden yapıldığında ortaya çıkan puan tablosudur."
    )
    bad = item_df[item_df["Kategori"] == "Kullanılmamalı"]["Madde"].tolist()
    good = [c for c in q_cols if c not in bad]
    new_pct = df[good].sum(axis=1).values / len(good) * 100
    old_pct = scores
    kr20_new = kr20(df[good]) if len(good) > 1 else 0

    sim = []
    for th in [40, 50, 60, 70]:
        op = (old_pct >= th).sum()
        np_ = (new_pct >= th).sum()
        sim.append(
            {
                "Eşik (%)": th,
                f"İlk analiz ({K} madde)": f"{op} (%{op / N * 100:.1f})",
                f"Kesme simülasyonu ({len(good)} madde)": f"{np_} (%{np_ / N * 100:.1f})",
                "Fark": np_ - op,
            }
        )
    st.dataframe(pd.DataFrame(sim), use_container_width=True)
    sc1, sc2, sc3 = st.columns(3)
    sc1.metric("Çıkarılan madde", len(bad))
    sc2.metric("KR-20 (mevcut)", f"{alpha:.3f}")
    sc3.metric(
        f"KR-20 ({len(good)} madde)",
        f"{kr20_new:.3f}",
        delta=f"{kr20_new - alpha:+.3f}",
    )

    # ---- Kritik Maddeler ----
    st.header("⚠️ Kritik Maddeler")
    neg = item_df[item_df["D"] < 0]
    if not neg.empty:
        st.error(
            f"**Negatif ayırt ediciliğe sahip {len(neg)} madde** — "
            f"yazan öğretim üyesi tarafından acil olarak incelenmelidir!"
        )
        st.dataframe(neg, use_container_width=True)

    ceil_ = item_df[item_df["p"] >= 0.95]
    if not ceil_.empty:
        with st.expander(f"🔝 Tavan: {len(ceil_)} madde (p≥0.95)"):
            st.dataframe(
                ceil_[["Madde", "p", "D", "Karar"]], use_container_width=True
            )

    flr_ = item_df[item_df["p"] <= 0.20]
    if not flr_.empty:
        with st.expander(f"🔻 Taban: {len(flr_)} madde (p≤0.20)"):
            st.dataframe(
                flr_[["Madde", "p", "D", "Karar"]], use_container_width=True
            )

    # ---- AI Değerlendirme ----
    ai_general = ""
    if api_key:
        import google.generativeai as genai

        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-3-flash-preview")
        st.header("🤖 Genel Değerlendirme")
        with st.spinner("Gemini analiz üretiyor..."):
            try:
                prompt = f"""Sen ölçme-değerlendirme uzmanısın. Aşağıdaki sınav verilerini analiz et.

KURALLAR:
- Türkçe akademik dilde yaz.
- İngilizce terimleri parantez içinde verme (skewness, kurtosis, ceiling effect gibi). Türkçe karşılıklarını kullan.
- "Tıp eğitimi standartları" yerine "ölçme değerlendirme ilkeleri" ifadesini kullan.
- "kanıtlamaktadır" yerine "göstermektedir" ifadesini tercih et.
- "vasfa sahip" yerine "özelliğe sahip" ifadesini kullan.
- "Güçlü yön" yerine "Dikkat Çekici Göstergeler" başlığı altında gerçekten dikkat çekici parametreleri belirt. Standart düzeyde olan şeyleri dikkat çekici olarak sunma. Dikkat çekici gösterge yoksa bu başlığı hiç açma.
- "Sorunlu Alanlar" yerine "Gelişime Açık Alanlar" başlığını kullan.
- Psikometrik açıdan gelişime açık alanları ve eksiklikleri net ve somut ifade et.
- Önerileri uygulanabilir ve spesifik yaz.
- Çeldirici kelimesini kullan (distraktör değil). Negatif ayırt edici ifadesini kullan (negatif D değil).
- Negatif ayırt edici maddeler için "İçerik uzmanı tarafından acil inceleme önerilir" yerine "soruyu yazan öğretim üyesi tarafından incelenmeli; soruda bilimsel bir hata veya anahtar hatası varsa iptal edilmeli, soru bankasında pasife alınmalı ya da düzeltilmelidir" ifadesini kullan.
- "Kullanılmamalı" maddelerin "ölçme hassasiyetine neredeyse hiçbir anlamlı katkı sağlamadığını" belirt (sadece sınavın verimliliğini düşürdüğünü değil).
- Çeldirici analizinde "çeldirici analizleri yapılmalı" yerine "çeldiricileri gözden geçirilmeli, hiçbir öğrencinin işaretlemediği işlemeyen çeldiriciler değiştirilerek maddenin ayırt ediciliği artırılmalıdır" ifadesini kullan.
- Madde havuzu revizyonunda "Kullanılmamalı" ve "Düzeltilmeli" kategorisindeki toplam madde sayısını ve yüzdesini belirt.
- Sınav boyutu optimizasyonunda "madde sayısı azaltılarak madde başına düşen süre artırılabilir" yerine "sınav setinde yer alacak soruların dikkatle hazırlanması gerektiği açıktır" ifadesini kullan.

VERİLER:
N={N}, K={K}, Ort={mn:.2f}±{sd:.2f} (100 üzerinden), Med={med:.1f}
KR-20={alpha:.3f}, SEM={sem:.2f}, Ferguson δ={fdelta:.3f}
Guttman={guttman:.3f}, Spearman-Brown={sb_corr:.3f}
Çarpıklık={skew:.2f}, Basıklık={kurt:.2f}
Zorluk: {item_df['Zorluk'].value_counts().to_dict()}
Ayırt edicilik: {cc.to_dict()}
Negatif ayırt edici madde: {(item_df['D'] < 0).sum()}
p≥0.95: {len(ceil_)}, p≤0.20: {len(flr_)}
Kaliteli alt küme ({len(qi)} madde) KR-20: {kr20_q if kr20_q else 'N/A'}
Kesme puanı sim: {len(bad)} madde çıkarılsa KR-20={kr20_new:.3f}
Tekrar eden soru: {tekrar_count}/{K} madde daha önceki sınavlardan tekrar kullanılmış{f" (Yeni: {yeni_count})" if tekrar_count > 0 else ""}

Başlıklar: 1. Gelişime Açık Alanlar  2. Dikkat Çekici Göstergeler (sadece gerçekten varsa)  3. Öneriler"""
                resp = model.generate_content(prompt)
                ai_general = resp.text
                st.markdown(ai_general)
            except Exception as e:
                st.warning(f"Gemini hatası: {e}")

    # ============================================================
    # DOCX REPORT
    # ============================================================
    st.header("📄 Rapor İndir")

    bufs = {}
    for name, fig in [("hist", fig1), ("violin", fig_v), ("scatter", fig2), ("matrix", fig_m)]:
        b = BytesIO()
        fig.savefig(b, format="png", dpi=180, bbox_inches="tight")
        b.seek(0)
        bufs[name] = b

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for lvl, sz, clr_hex in [(1, 16, "1F4E79"), (2, 13, "2E75B6"), (3, 11, "404040")]:
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "Arial"
        hs.font.size = Pt(sz)
        hs.font.color.rgb = RGBColor.from_string(clr_hex)

    # ---- Kapak ----
    for _ in range(6):
        doc.add_paragraph("")
    for text, size in [
        ("EGE ÜNİVERSİTESİ TIP FAKÜLTESİ", 18),
        ("TIP EĞİTİMİ ANABİLİM DALI", 14),
    ]:
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string("1F4E79")
    doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"{exam_name} ANALİZ RAPORU")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("1F4E79")
    for _ in range(4):
        doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run(f"Rapor Tarihi: {datetime.date.today().strftime('%d.%m.%Y')}").font.size = Pt(11)
    doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RAPORU HAZIRLAYAN")
    r.bold = True
    r.font.size = Pt(12)
    if analyst_name:
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(analyst_name)
        r.bold = True
        r.font.size = Pt(12)
    doc.add_page_break()

    # ---- Kısaltmalar ----
    doc.add_heading("Kısaltmalar ve Temel Kavramlar", 1)
    abbr = [
        ("p", "Güçlük İndeksi — doğru yanıt oranı (0.30–0.80 ideal)"),
        ("D", "Ayırt Edicilik İndeksi — üst-alt %27 grup farkı (≥0.30 iyi)"),
        ("r_pbi", "Düzeltilmiş nokta çift serili korelasyon (≥0.30 iyi)"),
        ("KR-20", "Kuder-Richardson 20 — iç tutarlılık güvenirliği (≥0.80 yüksek)"),
        ("SEM", "Standart Ölçme Hatası = SD × √(1−KR-20)"),
        ("Ferguson's δ", "Puanların yayılma genişliği (>0.90 iyi)"),
        ("Guttman", "Split-half güvenirlik katsayısı"),
        ("Çeldirici", "Yanlış seçenek; fonksiyonel ise ≥%5 seçilmelidir"),
        ("Negatif Ayırt Edici", "D<0: alt grup üst gruptan daha başarılı"),
    ]
    at = doc.add_table(rows=len(abbr) + 1, cols=2)
    at.style = "Light Grid Accent 1"
    at.rows[0].cells[0].text = "kısaltma ve kavram"
    at.rows[0].cells[1].text = "Açıklama"
    for cell in at.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (k, v) in enumerate(abbr):
        at.rows[i + 1].cells[0].text = k
        at.rows[i + 1].cells[1].text = v

    doc.add_heading("Not", 2)
    doc.add_paragraph(
        "Bu rapordaki madde analizleri Klasik Test Teorisi (KTT) çerçevesinde, "
        "üst-alt %27 grup yöntemiyle yapılmıştır. Kelley (1939) ve Ebel (1965) "
        "tarafından gösterildiği üzere bu oran ayırt edicilik indeksi için en güçlü "
        "istatistiksel kestirim performansını vermektedir. Puanlar soru sayısından "
        "bağımsız olarak 100 üzerinden normalize edilmiştir."
    )

    # ---- Bölüm 1: Genel ----
    doc.add_page_break()
    doc.add_heading("1. Genel Psikometrik Göstergeler", 1)
    sd_ = [
        ("Gösterge", "Değer", "Yorum"),
        ("Öğrenci Sayısı (N)", str(N), ""),
        ("Madde Sayısı (sınav setindeki soru sayısı)", str(K), ""),
        ("Ortalama ± SD", f"{mn:.2f} ± {sd:.2f}", "100 üzerinden"),
        ("Ortanca", f"{med:.1f}", ""),
        ("Q1 – Q3", f"{q1v:.1f} – {q3v:.1f}", f"IQR = {q3v - q1v:.1f}"),
        ("Min – Max", f"{scores.min():.0f} – {scores.max():.0f}", ""),
        ("Çarpıklık / Basıklık", f"{skew:.3f} / {kurt:.3f}", "Sola çarpık" if skew < -0.5 else "Normal"),
        ("≥60 puan (Başarılı öğrenci oranı)", f"{p60} (%{p60 / N * 100:.1f})", ""),
        ("<40 puan (Başarısız öğrenci oranı)", f"{f40} (%{f40 / N * 100:.1f})", ""),
        ("KR-20 (tüm maddeler)", f"{alpha:.3f}", "Yüksek" if alpha >= 0.80 else "Kabul edilebilir"),
        (f"KR-20 (kaliteli {len(qi)} madde)", f"{kr20_q:.3f}" if kr20_q else "—", "Önerilen/KE ∩ Mükemmel/İyi"),
        ("Ferguson's δ", f"{fdelta:.3f}", "İyi" if fdelta >= 0.90 else "Düşük"),
        ("Guttman Split-Half", f"{guttman:.3f}", ""),
        ("Spearman-Brown", f"{sb_corr:.3f}", ""),
        ("SEM", f"{sem:.2f}", ""),
        ("%95 Güven Aralığı", f"±{ci95:.1f} puan", ""),
        ("Tekrar Eden Soru", f"{tekrar_count}/{K} (%{tekrar_count/K*100:.1f})",
         f"Yeni: {yeni_count}" if tekrar_count > 0 else "Tüm sorular yeni"),
    ]
    gt = doc.add_table(rows=len(sd_), cols=3)
    gt.style = "Light Grid Accent 1"
    for i, (g, d, y) in enumerate(sd_):
        gt.rows[i].cells[0].text = g
        gt.rows[i].cells[1].text = d
        gt.rows[i].cells[2].text = y
        if i == 0:
            for cell in gt.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    # ---- Bölüm 2: Puan Dağılımı ----
    doc.add_heading("2. Öğrencilerin Puan Dağılımı", 1)
    doc.add_picture(bufs["hist"], width=Inches(5.8))
    doc.add_paragraph("")
    doc.add_picture(bufs["violin"], width=Inches(5.8))

    # ---- Bölüm 3: Madde Analizi ----
    doc.add_page_break()
    doc.add_heading("3. Madde Analizi Detay Tablosu", 1)

    # Özet satırı — bold ve kırmızı
    summary_para = doc.add_paragraph()
    summary_text = (
        f"Mükemmel: {cc.get('Mükemmel', 0)}  ·  İyi: {cc.get('İyi', 0)}  ·  "
        f"Düzeltilmeli: {cc.get('Düzeltilmeli', 0)}  ·  "
        f"Kullanılmamalı: {cc.get('Kullanılmamalı', 0)}  ·  "
        f"Negatif Ayırt Edici: {(item_df['D'] < 0).sum()}"
    )
    r = summary_para.add_run(summary_text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Renk açıklaması (Bölüm 4 ile tutarlı)
    doc.add_paragraph(
        "Tablodaki satır renkleri maddelerin ayırt edicilik kategorisine göre belirlenmiştir. "
        "Buna göre yeşil = mükemmel/iyi (sakla), sarı = düzeltilmeli (gözden geçir), "
        "turuncu = kullanılmamalı (revize et) ve kırmızı = negatif ayırt edici (sınav setinden çıkar) "
        "anlamına gelmektedir."
    )

    hdrs = ["Madde", "p", "Zorluk", "D", "r_pbi", "Kategori", "Çeldirici\n(işlevsiz/toplam)", "Önceki Kullanım", "Karar"]
    mt = doc.add_table(rows=len(item_df) + 1, cols=len(hdrs))
    mt.style = "Light Grid Accent 1"
    for j, h in enumerate(hdrs):
        mt.rows[0].cells[j].text = h
        for run in mt.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    # Kategori → arka plan renk eşleştirmesi (Karar Destek Matrisi ile tutarlı)
    cat_colors = {
        "Mükemmel": "C8E6C9",    # yeşil
        "İyi": "C8E6C9",         # yeşil
        "Düzeltilmeli": "FFF9C4", # sarı
        "Kullanılmamalı": "FFE0B2", # turuncu
    }

    for i, (_, row) in enumerate(item_df.iterrows()):
        cols = ["Madde", "p", "Zorluk", "D", "r_pbi", "Kategori", "Çeldirici", "Önceki Kullanım", "Karar"]
        kategori = row["Kategori"]
        d_val = row["D"]

        # Negatif ayırt edici → kırmızı
        if d_val < 0:
            bg_color = "FFCDD2"
        else:
            bg_color = cat_colors.get(kategori, None)

        for j, col in enumerate(cols):
            cell = mt.rows[i + 1].cells[j]
            cell.text = str(row[col])
            # Arka plan rengi uygula
            if bg_color:
                from docx.oxml.parser import parse_xml
                from docx.oxml.ns import nsdecls
                shading_elm = parse_xml(
                    f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color}"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading_elm)

    # ---- Bölüm 4: Karar Destek Matrisi ----
    doc.add_heading("4. Karar Destek Matrisi", 1)
    doc.add_paragraph(
        "Aşağıdaki tabloda sınav setine ilişkin karar destek matrisi sunulmaktadır. "
        "Tablodaki renkler maddelerin ayırt edicilik ve zorluk indeksine göre dağılımını "
        "ifade etmektedir. Buna göre yeşil = sakla, sarı = gözden geçir, turuncu = revize et "
        "ve kırmızı = sınav setinden çıkar anlamına gelmektedir."
    )
    doc.add_picture(bufs["matrix"], width=Inches(5.5))

    # ---- Bölüm 5: Scatter ----
    doc.add_heading("5. Zorluk ve Ayırt Edicilik İndekslerinin Birlikteliği", 1)
    doc.add_picture(bufs["scatter"], width=Inches(5.5))

    # ---- Bölüm 6: Kesme Puanı ----
    doc.add_page_break()
    doc.add_heading("6. Kesme Puanı Simülasyonu", 1)
    doc.add_paragraph(
        "Kesme puanı simülasyonu; sınav analizleri sonrasında sınav setindeki kullanılmamalı "
        "kategorisindeki sorular dikkate alınmadan analizler yeniden yapıldığında ortaya çıkan "
        f"puan tablosudur. Buna göre bu sınavda {len(bad)} adet 'Kullanılmamalı' kategorisindeki "
        "madde çıkarıldığında puanlar 100 üzerinden yeniden hesaplandığında aşağıdaki tablo "
        "ortaya çıkmıştır."
    )
    doc.add_paragraph(f"KR-20: {alpha:.3f} → {kr20_new:.3f} ({kr20_new - alpha:+.3f})")
    st_ = doc.add_table(rows=len(sim) + 1, cols=4)
    st_.style = "Light Grid Accent 1"
    sim_hdrs = ["Eşik (%)", f"İlk analiz ({K} madde)", f"Kesme simülasyonu ({len(good)} madde)", "Fark"]
    for j, h in enumerate(sim_hdrs):
        st_.rows[0].cells[j].text = h
        for run in st_.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True
    for i, s in enumerate(sim):
        st_.rows[i + 1].cells[0].text = str(s["Eşik (%)"])
        st_.rows[i + 1].cells[1].text = str(s[f"İlk analiz ({K} madde)"])
        st_.rows[i + 1].cells[2].text = str(s[f"Kesme simülasyonu ({len(good)} madde)"])
        st_.rows[i + 1].cells[3].text = str(s["Fark"])

    # ---- Bölüm 7: Kritik Maddeler ----
    if not neg.empty:
        doc.add_heading("7. Kritik Maddeler — Negatif Ayırt Edici", 1)
        doc.add_paragraph(
            "Sınav analizi sonucunda maddelerde alt gruptaki öğrenciler üst gruptan daha yüksek "
            "doğru yanıt oranına sahip olması Negatif Ayırt edicilik olarak tanımlanmaktadır. "
            "Bu durumun saptandığı sorularda hatalı cevap anahtarı, soru kökünde belirsizlik "
            "veya çoklu doğru yanıt ihtimali dikkate alınmalıdır. Negatif Ayırt edicilik "
            "belirlenen soru; yazan öğretim üyesi tarafından acil olarak incelemelidir."
        )
        doc.add_paragraph(
            f"Bu sınavda Negatif Ayırt edicilik saptanan sorular;"
        )
        doc.add_paragraph(f"Maddeler: {', '.join(neg['Madde'].tolist())}")

    # ---- Bölüm 8: AI ----
    if ai_general:
        doc.add_page_break()
        next_section = 8 if not neg.empty else 7
        doc.add_heading(f"{next_section}. Genel Değerlendirme*", 1)
        intro = doc.add_paragraph(
            "Ölçme-değerlendirme ilkeleri doğrultusunda, paylaşılan sınav verilerinin "
            "psikometrik analizine ilişkin değerlendirmeler aşağıda sunulmuştur."
        )
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for line in ai_general.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("#"):
                # Başlık — heading 2, bold zaten ama ek olarak garantiliyoruz
                clean_title = re.sub(r"\*\*(.+?)\*\*", r"\1", line.lstrip("# "))
                h = doc.add_heading(clean_title, level=2)
                for run in h.runs:
                    run.bold = True
            elif line.startswith("- ") or line.startswith("* "):
                # Bullet item — bold kalıpları koru, iki yana yasla
                raw_text = line[2:]
                bp = doc.add_paragraph(style="List Bullet")
                bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # **bold** kalıplarını parse et
                parts = re.split(r"(\*\*.+?\*\*)", raw_text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = bp.add_run(part[2:-2])
                        r.bold = True
                    else:
                        bp.add_run(part)
            else:
                # Normal paragraf — bold kalıpları koru, iki yana yasla
                pp = doc.add_paragraph()
                pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parts = re.split(r"(\*\*.+?\*\*)", line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = pp.add_run(part[2:-2])
                        r.bold = True
                    else:
                        pp.add_run(part)

    # ---- AI Footnote ----
    if ai_general:
        doc.add_paragraph("")
        fn = doc.add_paragraph()
        r = fn.add_run(
            '*Bu raporun oluşturulmasında yapay zekadan destek alınmış olup '
            'genel değerlendirme ifadeleri Tıp Eğitimi Danışman Öğretim Üyesi '
            'tarafından hazırlanmıştır.'
        )
        r.italic = True
        r.font.size = Pt(9)

    # ---- Son Sayfa: İmza ----
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph("")
    for text in ["Ege Üniversitesi Tıp Fakültesi", "Tıp Eğitimi Anabilim Dalı"]:
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.add_run(text).bold = True
    doc.add_paragraph("")

    if analyst_name:
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run("Tıp Eğitimi Danışman Öğretim Üyesi")
        r.bold = True
        r.font.size = Pt(11)
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(analyst_name)
        r.font.size = Pt(12)
        r.bold = True

    doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.add_run(
        f"Rapor Tarihi: {datetime.date.today().strftime('%d.%m.%Y')}"
    ).font.size = Pt(10)

    # Save
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        "📥 Tam Rapor İndir (.docx)",
        data=buf,
        file_name=f"{exam_name} Analiz Raporu.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True,
    )

elif not run_btn:
    st.info(
        "👆 İki Excel dosyasını yükleyin, danışman adını girin ve "
        "**Analizi Başlat** butonuna basın."
    )
